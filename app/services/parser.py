from __future__ import annotations

from io import BytesIO
from pathlib import Path


class ParseError(Exception):
    pass


def extract_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _pdf_text(data)
    if suffix in {".docx", ".doc"}:
        if suffix == ".doc":
            raise ParseError("Legacy .doc is not supported. Save as .docx, .pdf, or .txt.")
        return _docx_text(data)
    if suffix in {".txt", ".md"}:
        return data.decode("utf-8", errors="replace")
    raise ParseError("Upload a PDF, DOCX, or TXT file.")


def _pdf_text(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(data))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    text = "\n".join(pages).strip()
    if not text:
        raise ParseError("No extractable text in that PDF. Try a text-based export, not a scan.")
    return text


def _docx_text(data: bytes) -> str:
    from docx import Document

    document = Document(BytesIO(data))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" · ".join(cells))
    text = "\n".join(parts).strip()
    if not text:
        raise ParseError("No extractable text in that Word file.")
    return text
